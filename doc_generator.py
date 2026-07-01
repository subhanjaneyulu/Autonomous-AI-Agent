import os
from datetime import datetime

from docx import Document
from docx.shared import Pt


def generate_docx(document_content: str, document_type: str) -> str:
    """
    Generate a Microsoft Word (.docx) document.

    Args:
        document_content (str): Generated business document.
        document_type (str): Type of the document.

    Returns:
        str: Path of the generated document.
    """

    # Create output folder
    os.makedirs("output", exist_ok=True)

    # Create document
    doc = Document()

    # Title
    title = doc.add_heading(document_type, level=1)
    title.runs[0].font.size = Pt(20)

    # Add generated content
    for line in document_content.split("\n"):

        line = line.strip()

        if not line:
            continue

        # Headings
        if line.isupper():
            doc.add_heading(line.title(), level=2)

        # Bullet Points
        elif line.startswith("-"):
            doc.add_paragraph(line, style="List Bullet")

        else:
            paragraph = doc.add_paragraph()
            paragraph.add_run(line).font.size = Pt(11)

    # File Name
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")

    filename = f"{document_type.replace(' ', '_')}_{timestamp}.docx"

    filepath = os.path.join("output", filename)

    # Save document
    doc.save(filepath)

    return filepath